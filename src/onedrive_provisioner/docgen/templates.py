"""License-to-section template registry.

Each entry maps a license keyword pattern to a section builder function.
Builders append headings, numbered / bulleted steps, hyperlinks, and
screenshots into the document, matching the reference guide format.

To add a new service, add an entry to ``LICENSE_SECTION_REGISTRY``.
"""
from __future__ import annotations

import os
import re
from typing import Callable, Dict, List, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Type alias for a section builder
SectionBuilder = Callable  # (doc: Document, screenshots_dir: str) -> None

# ── Helpers ──

def _add_screenshot(doc: Document, screenshots_dir: str, filename: str,
                    width: float = 5.0) -> None:
    """Insert an image if it exists, otherwise silently skip."""
    path = os.path.join(screenshots_dir, filename)
    if os.path.isfile(path):
        doc.add_picture(path, width=Inches(width))


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _bullet(doc: Document, text: str) -> None:
    """Add a bullet-style list paragraph."""
    doc.add_paragraph(text, style="List Bullet")


def _bullet_with_link(doc: Document, text_before: str, url: str,
                      link_text: str, text_after: str = "") -> None:
    """Add a bullet paragraph that contains an inline hyperlink."""
    p = doc.add_paragraph(style="List Bullet")
    if text_before:
        p.add_run(text_before)
    p.add_run(" \U0001f449 ")  # 👉
    _add_hyperlink(p, url, link_text)
    if text_after:
        p.add_run(text_after)


# ── Section builders ──

def _section_copilot_studio(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Copilot Studio", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://copilotstudio.microsoft.com",
                      "Copilot Studio")
    _bullet(doc, "Sign in using the same credentials.")
    _add_screenshot(doc, screenshots_dir, "copilot_studio_home.png", 5.5)


def _section_m365(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access the M365 Portal", level=2)
    _bullet_with_link(doc,
                      "Go to the portal using the following link:",
                      "https://portal.office.com",
                      "M365 Portal")
    _bullet(doc, "Log in with the credentials provided by your trainer.")
    _add_screenshot(doc, screenshots_dir, "m365_login.png", 3.4)
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _add_screenshot(doc, screenshots_dir, "stay_signed_in.png", 3.4)
    _bullet(doc, "Upon login you will see the M365 home page that you will be "
                 "using throughout the hack.")
    _add_screenshot(doc, screenshots_dir, "m365_copilot_chat.png", 5.5)


def _section_power_bi(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Power BI", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://app.powerbi.com",
                      "Power BI")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _add_screenshot(doc, screenshots_dir, "stay_signed_in.png", 3.4)
    _add_screenshot(doc, screenshots_dir, "powerbi_home.png", 5.0)


def _section_power_apps(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Power Apps", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://make.powerapps.com",
                      "Power Apps")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _add_screenshot(doc, screenshots_dir, "powerapps_home.png", 5.0)


def _section_power_automate(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Power Automate", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://make.powerautomate.com",
                      "Power Automate")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _add_screenshot(doc, screenshots_dir, "powerautomate_home.png", 5.0)


def _section_teams(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Microsoft Teams", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://teams.microsoft.com",
                      "Microsoft Teams")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _bullet(doc, "You can also download the Teams desktop or mobile app "
                 "for a better experience.")
    _add_screenshot(doc, screenshots_dir, "teams_home.png", 5.0)


def _section_azure_portal(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access the Azure Portal", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://portal.azure.com",
                      "Azure Portal")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _bullet(doc, "Once logged in you will see that you have been assigned an Azure "
                 "subscription. You will be using the same subscription throughout the hack.")
    _bullet(doc, "Navigate to your assigned resource group to view available resources.")
    _bullet(doc, "Please make sure that you stop resources/services when not in use "
                 "to save Azure cost as all subscriptions come with a budget limit.")
    _add_screenshot(doc, screenshots_dir, "azure_portal.png", 5.0)


def _section_azure_ai_foundry(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Azure AI Foundry", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://ai.azure.com",
                      "Azure AI Foundry")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _bullet(doc, "Explore AI models, playgrounds, and deployment options.")
    _add_screenshot(doc, screenshots_dir, "ai_foundry.png", 5.0)


def _section_github(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access GitHub", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://github.com",
                      "GitHub")
    _bullet(doc, "Sign in with credentials provided separately by your hack admin.")
    _bullet(doc, "Navigate to the repository shared for this hackathon.")
    _add_screenshot(doc, screenshots_dir, "github_home.png", 5.0)


def _section_windows365(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Windows 365", level=2)
    _bullet_with_link(doc,
                      "Navigate to Windows 365 and enter the credentials that you have "
                      "received for this hack:",
                      "https://windows365.microsoft.com",
                      "Windows 365")
    _bullet(doc, 'Once logged in you will see the "Your Windows experience, wherever '
                 'you are" pop-up. Click Next.')
    _add_screenshot(doc, screenshots_dir, "w365_welcome.png", 5.0)
    _bullet(doc, 'On the "Enhance your experience with a guided tour" page, click Not now.')
    _add_screenshot(doc, screenshots_dir, "w365_guided_tour.png", 5.0)
    _bullet(doc, "You have successfully signed in to Windows 365. Now, click on the Cloud PC.")
    _add_screenshot(doc, screenshots_dir, "w365_cloud_pc.png", 5.0)
    _bullet(doc, "A new browser tab will open. In the In-session settings pane, "
                 "keep the default settings and click Connect.")
    _add_screenshot(doc, screenshots_dir, "w365_connect.png", 5.0)
    _bullet(doc, "Provide the password and click Sign in.")
    _add_screenshot(doc, screenshots_dir, "w365_signin.png", 5.0)
    _bullet(doc, "You have successfully logged in to the Cloud PC.")
    _add_screenshot(doc, screenshots_dir, "w365_desktop.png", 5.0)


def _section_devops(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Azure DevOps", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://dev.azure.com",
                      "Azure DevOps")
    _bullet(doc, "Sign in using the same credentials.")
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _add_screenshot(doc, screenshots_dir, "devops_home.png", 5.0)


# ── Registry: license keyword pattern -> (section_title, builder) ──
# Order determines the order sections appear in the document.

LICENSE_SECTION_REGISTRY: List[Tuple[re.Pattern, str, SectionBuilder]] = [
    (re.compile(r"COPILOT_STUDIO|COPILOTSTUDIO|COPILOT.*STUDIO", re.I),
     "Copilot Studio", _section_copilot_studio),

    (re.compile(r"M365|O365|OFFICE365|ENTERPRISE|SPE_E|SPE_F|DEVELOPERPACK|"
                r"STANDARDPACK|DESKLESSPACK|E3|E5|G3|G5", re.I),
     "Microsoft 365", _section_m365),

    (re.compile(r"POWER_BI|PBI_|BI_AZURE", re.I),
     "Power BI", _section_power_bi),

    (re.compile(r"POWERAPPS|POWER_APPS", re.I),
     "Power Apps", _section_power_apps),

    (re.compile(r"FLOW_|POWERAUTOMATE|POWER_AUTOMATE", re.I),
     "Power Automate", _section_power_automate),

    (re.compile(r"TEAMS_ESSENTIALS|TEAMS_EXPLORATORY|TEAMS_COMMERCIAL", re.I),
     "Microsoft Teams", _section_teams),

    (re.compile(r"AZURE|AZR_", re.I),
     "Azure Portal", _section_azure_portal),

    (re.compile(r"AI_FOUNDRY|COGNITIVE|OPENAI", re.I),
     "Azure AI Foundry", _section_azure_ai_foundry),

    (re.compile(r"GITHUB", re.I),
     "GitHub", _section_github),

    (re.compile(r"WINDOWS_365|CPC_B_|W365|CLOUD_PC", re.I),
     "Windows 365", _section_windows365),

    (re.compile(r"DEVOPS|AZURE_DEVOPS", re.I),
     "Azure DevOps", _section_devops),
]


# ── GCC-specific section builders ──

def _section_copilot_studio_gcc(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access Copilot Studio (GCC)", level=2)
    _bullet_with_link(doc,
                      "Open the following link:",
                      "https://gcc.powerva.microsoft.us/",
                      "Copilot Studio (GCC)")
    _bullet(doc, "Sign in using the same credentials.")
    _add_screenshot(doc, screenshots_dir, "copilot_studio_home.png", 5.5)


def _section_m365_gcc(doc: Document, screenshots_dir: str) -> None:
    doc.add_heading("Access the M365 Portal (GCC)", level=2)
    _bullet_with_link(doc,
                      "Go to the portal using the following link:",
                      "https://portal.office365.us",
                      "M365 Portal (GCC)")
    _bullet(doc, "Log in with the credentials provided by your trainer.")
    _add_screenshot(doc, screenshots_dir, "m365_login.png", 3.4)
    _bullet(doc, 'On the "Stay signed in" page, click Yes.')
    _add_screenshot(doc, screenshots_dir, "stay_signed_in.png", 3.4)
    _bullet(doc, "Upon login you will see M365 Copilot window that you will be "
                 "using throughout the hack.")
    _add_screenshot(doc, screenshots_dir, "m365_copilot_chat.png", 5.5)


GCC_SECTION_REGISTRY: List[Tuple[re.Pattern, str, SectionBuilder]] = [
    (re.compile(r"COPILOT_STUDIO|COPILOTSTUDIO|COPILOT.*STUDIO", re.I),
     "Copilot Studio (GCC)", _section_copilot_studio_gcc),

    (re.compile(r"M365|O365|OFFICE365|ENTERPRISE|SPE_E|SPE_F|DEVELOPERPACK|"
                r"STANDARDPACK|DESKLESSPACK|E3|E5|G3|G5", re.I),
     "Microsoft 365 (GCC)", _section_m365_gcc),
]


# ── License descriptions for the "Licenses" summary section ──

LICENSE_DESCRIPTIONS: Dict[str, str] = {
    "COPILOT_STUDIO": "Build, customize, and deploy AI copilots using low-code tools with enterprise integrations.",
    "CCIBOTS_PRIVPREV_VIRAL": "Build, customize, and deploy AI copilots using low-code tools with enterprise integrations.",
    "Microsoft_Copilot_Studio_User": "Build, customize, and deploy AI copilots using low-code tools with enterprise integrations.",
    "O365_BUSINESS_PREMIUM": "Cloud-based productivity suite with email, Teams, and Office apps (web/mobile).",
    "SPB": "Cloud-based productivity suite with email, Teams, and Office apps (web/mobile).",
    "Microsoft_365_Copilot": "AI assistant integrated across Microsoft 365 apps to automate tasks and enhance productivity.",
    "COPILOT_FOR_MICROSOFT_365": "AI assistant integrated across Microsoft 365 apps to automate tasks and enhance productivity.",
    "M365_COPILOT": "AI assistant integrated across Microsoft 365 apps to automate tasks and enhance productivity.",
    "SPE_E3": "Enterprise productivity suite with advanced compliance, security, and management features.",
    "ENTERPRISEPACK": "Enterprise cloud productivity suite with email, Teams, and Office apps.",
    "SPE_E5": "Premium enterprise suite with advanced analytics, security, and voice capabilities.",
    "ENTERPRISEPREMIUM": "Premium enterprise suite with advanced analytics, compliance, and voice capabilities.",
    "POWER_BI_PRO": "Business analytics platform for interactive visualizations and business intelligence.",
    "PBI_PREMIUM_PER_USER": "Premium analytics with advanced AI, dataflows, and paginated reports per user.",
    "POWERAPPS_PER_USER": "Build custom business apps with low-code tools and enterprise-grade data connectors.",
    "POWER_APPS_PER_USER": "Build custom business apps with low-code tools and enterprise-grade data connectors.",
    "TEAMS_ESSENTIALS": "Core Microsoft Teams collaboration: chat, video meetings, and file sharing.",
    "Teams_Ess": "Core Microsoft Teams collaboration: chat, video meetings, and file sharing.",
    "TEAMS_PREMIUM": "Enhanced Teams with AI-powered meetings, webinars, and advanced security.",
    "Teams_Premium": "Enhanced Teams with AI-powered meetings, webinars, and advanced security.",
    "FLOW_FREE": "Automate workflows across apps and services with Power Automate.",
    "WINDOWS_365": "Dedicated Cloud PC with full Windows OS. Run apps, development tools, secure workspace anywhere.",
    "CPC_B_2C_8RAM_256GB": "Dedicated Cloud PC with full Windows OS. Run apps, development tools, secure workspace anywhere.",
    "M365_G3_GOV": "Government-grade M365 with compliance, security, and productivity tools.",
    "M365_G5_GOV": "Premium government M365 with advanced analytics, compliance, and security.",
}


def get_license_description(sku: str) -> str:
    """Return a human-friendly description for a license SKU."""
    if not sku:
        return ""
    # Direct match
    if sku in LICENSE_DESCRIPTIONS:
        return LICENSE_DESCRIPTIONS[sku]
    # Case-insensitive search
    upper = sku.upper()
    for key, desc in LICENSE_DESCRIPTIONS.items():
        if key.upper() == upper:
            return desc
    # Partial match
    for key, desc in LICENSE_DESCRIPTIONS.items():
        if key.upper() in upper or upper in key.upper():
            return desc
    return ""


def get_sections_for_licenses(
    license_skus: List[str],
    *,
    is_gcc: bool = False,
) -> List[Tuple[str, SectionBuilder]]:
    """Return matching section builders for the given license SKU part numbers.

    When ``is_gcc`` is True, GCC-specific variant sections (e.g. Copilot Studio GCC,
    M365 GCC) are appended after the standard sections.
    """
    seen: set = set()
    sections: List[Tuple[str, SectionBuilder]] = []
    all_skus = " ".join(license_skus)
    for pattern, title, builder in LICENSE_SECTION_REGISTRY:
        if title in seen:
            continue
        if pattern.search(all_skus):
            seen.add(title)
            sections.append((title, builder))

    if is_gcc:
        for pattern, title, builder in GCC_SECTION_REGISTRY:
            if title in seen:
                continue
            if pattern.search(all_skus):
                seen.add(title)
                sections.append((title, builder))

    return sections
