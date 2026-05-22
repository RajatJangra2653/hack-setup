"""Core document generator — builds a professional .docx Admin/Trainer Guide.

Matches the San Francisco reference guide format:
- A4 page, 0.5" margins
- Calibri body, Heading 2 = 12pt teal (#0F4761)
- Blue (#0070C0) centered section headers
- Normal Table style tables with bold header rows
- Bullet list steps with 👉 hyperlinks and inline screenshots
"""
from __future__ import annotations

import io
import os
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from onedrive_provisioner.entra.models import license_display_name

from .templates import get_sections_for_licenses, get_license_description

# Default screenshots directory (relative to project root)
_DEFAULT_SCREENSHOTS = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))),
    "assets", "screenshots",
)

# Colors matching the reference document
BLUE = RGBColor(0x00, 0x70, 0xC0)
TEAL = RGBColor(0x0F, 0x47, 0x61)
GRAY = RGBColor(0x88, 0x88, 0x88)


class DocGenerator:
    """Generates a Trainer/Admin Guide .docx from hack state data."""

    def __init__(self, screenshots_dir: Optional[str] = None) -> None:
        self._screenshots_dir = screenshots_dir or _DEFAULT_SCREENSHOTS

    def generate(self, state: Dict[str, Any]) -> bytes:
        """Generate a .docx document from hack state and return raw bytes."""
        doc = Document()
        self._setup_page(doc)
        self._setup_styles(doc)

        self._add_title(doc, state)
        self._add_intro(doc, state)
        self._add_section_header(doc, "Environment Access and Licensing Summary")
        self._add_user_access_structure(doc, state)
        self._add_license_allocation(doc, state)
        self._add_budget_allocation(doc, state)
        self._add_trainer_note(doc)
        self._add_license_descriptions(doc, state)
        self._add_section_header(doc, "Access Guide: Overview & Login Instructions")
        self._add_access_intro(doc)
        self._add_access_sections(doc, state)
        self._add_support_contact(doc, state)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def get_filename(self, state: Dict[str, Any]) -> str:
        """Return the recommended filename for the document."""
        hack_name = state.get("hackName") or state.get("prefix", "hack")
        safe = "".join(c if c.isalnum() or c in "-_ " else "" for c in hack_name).strip()
        safe = safe.replace(" ", "-") or "hack"
        return f"{safe}-Admin-Guide.docx"

    # ──────────── Page setup ────────────

    @staticmethod
    def _setup_page(doc: Document) -> None:
        """A4 page, 0.5 inch margins on all sides (matching reference)."""
        section = doc.sections[0]
        section.page_width = Inches(8.27)   # A4 width
        section.page_height = Inches(11.69)  # A4 height
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    @staticmethod
    def _setup_styles(doc: Document) -> None:
        """Set up default fonts and heading styles to match reference."""
        # Body style
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # Heading 2 = 12pt teal, matching reference's Heading 2
        h2 = doc.styles["Heading 2"]
        h2.font.name = "Calibri"
        h2.font.size = Pt(16)
        h2.font.color.rgb = TEAL

    # ──────────── Title ────────────

    @staticmethod
    def _add_title(doc: Document, state: Dict[str, Any]) -> None:
        """Centered blue bold title matching reference format."""
        hack_name = state.get("hackName") or state.get("prefix", "Hackathon")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{hack_name} - Admin Guide: Overview & Login Instructions")
        run.bold = True
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"

    # ──────────── Intro ────────────

    @staticmethod
    def _add_intro(doc: Document, state: Dict[str, Any]) -> None:
        doc.add_paragraph(
            "This document provides a comprehensive overview of the hackathon "
            "environment, including instructions on how to use the provisioned "
            "resources and services. Please review this guide carefully before "
            "the event."
        )

    # ──────────── Section header (blue centered bold) ────────────

    @staticmethod
    def _add_section_header(doc: Document, text: str) -> None:
        """Centered blue bold section divider, matching reference style."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"

    # ──────────── User & Access Structure ────────────

    @staticmethod
    def _add_user_access_structure(doc: Document, state: Dict[str, Any]) -> None:
        doc.add_heading("User & Access Structure", level=2)

        users = state.get("users", [])
        total = len(users)
        admins = sum(1 for u in users if u.get("isAdmin"))
        groups = state.get("groups", [])
        config = state.get("config", {})
        teams = config.get("teams", len(groups))
        users_per_team = config.get("usersPerTeam", 0)
        mode = config.get("mode", state.get("mode", "team"))

        if mode == "team" and teams:
            doc.add_paragraph(
                f"A total of {total} users are organized into {teams} teams"
                + (f", with {users_per_team} users per team." if users_per_team else ".")
            , style="List Bullet")
            # Subscription mapping
            sub_ids = state.get("subscriptionIds") or config.get("subscriptionIds") or []
            if isinstance(sub_ids, str):
                sub_ids = [s.strip() for s in sub_ids.split(",") if s.strip()]
            if sub_ids:
                doc.add_paragraph(
                    "Each team has access to one dedicated Azure subscription.",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph(
                f"A total of {total} users have been provisioned in flat mode.",
                style="List Bullet",
            )

        if admins:
            sub_ids = state.get("subscriptionIds") or config.get("subscriptionIds") or []
            if isinstance(sub_ids, str):
                sub_ids = [s.strip() for s in sub_ids.split(",") if s.strip()]
            sub_count = len(sub_ids) if sub_ids else teams
            suffix = f" access across all {sub_count} subscriptions." if sub_count else "."
            doc.add_paragraph(
                f"Additionally, {admins} admin{'s have' if admins > 1 else ' has'}"
                + suffix,
                style="List Bullet",
            )

        # Subscription table
        sub_ids = state.get("subscriptionIds") or config.get("subscriptionIds") or []
        if isinstance(sub_ids, str):
            sub_ids = [s.strip() for s in sub_ids.split(",") if s.strip()]
        sub_names = state.get("subscriptionNames") or config.get("subscriptionNames") or []
        if sub_ids:
            doc.add_paragraph(
                "Subscriptions in Use:",
                style="List Bullet",
            )
            table = doc.add_table(rows=len(sub_ids), cols=1)
            table.style = "Table Grid"
            for i, sid in enumerate(sub_ids):
                name = sub_names[i] if i < len(sub_names) else sid
                table.rows[i].cells[0].text = name

    # ──────────── License Allocation ────────────

    @staticmethod
    def _add_license_allocation(doc: Document, state: Dict[str, Any]) -> None:
        doc.add_heading("License Allocation", level=2)

        users = state.get("users", [])
        license_counts: Dict[str, int] = {}
        for u in users:
            for lic in u.get("licenses", []):
                license_counts[lic] = license_counts.get(lic, 0) + 1

        if not license_counts:
            doc.add_paragraph("No licenses assigned.")
            return

        # Simple table with bold header row (matching reference)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "License Type"
        hdr[1].text = "Number of Licenses"
        # Bold the header row
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for lic, count in sorted(license_counts.items()):
            row = table.add_row()
            row.cells[0].text = license_display_name(lic)
            row.cells[1].text = str(count)

    # ──────────── Azure Budget Allocation ────────────

    @staticmethod
    def _add_budget_allocation(doc: Document, state: Dict[str, Any]) -> None:
        config = state.get("config", {})
        budget_per_team = config.get("budgetPerTeam") or state.get("budgetPerTeam")
        total_budget = config.get("totalBudget") or state.get("totalBudget")
        currency = config.get("budgetCurrency") or state.get("budgetCurrency") or "USD"

        # Try to compute total from per-team * teams
        teams = int(config.get("teams", 0))
        if budget_per_team and not total_budget and teams:
            try:
                total_budget = float(budget_per_team) * teams
            except (ValueError, TypeError):
                pass

        if not budget_per_team and not total_budget:
            return  # no budget info — skip section entirely

        doc.add_heading("Azure Budget Allocation", level=2)

        sym = "$" if currency.upper() == "USD" else f"{currency} "

        if budget_per_team:
            doc.add_paragraph(
                f"{sym}{budget_per_team} per team",
                style="List Bullet",
            )
        if total_budget:
            doc.add_paragraph(
                f"Total Azure budget: {sym}{total_budget}",
            )

    # ──────────── License Descriptions ────────────

    @staticmethod
    def _add_license_descriptions(doc: Document, state: Dict[str, Any]) -> None:
        """Add a 'Licenses' section with short description for each assigned license."""
        users = state.get("users", [])
        sku_set: set = set()
        for u in users:
            for lic in u.get("licenses", []):
                sku_set.add(lic)

        if not sku_set:
            return

        doc.add_heading("Licenses", level=2)

        for sku in sorted(sku_set):
            friendly = license_display_name(sku)
            desc = get_license_description(sku)
            p = doc.add_paragraph(style="List Bullet")
            run_name = p.add_run(friendly)
            run_name.bold = True
            if desc:
                p.add_run(f"\n{desc}")

    # ──────────── Trainer Note ────────────

    @staticmethod
    def _add_trainer_note(doc: Document) -> None:
        """Blue highlighted note for trainers/admins, matching reference."""
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Note for the Trainers / Admins: ")
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"

        doc.add_paragraph(
            "It is important that you test the hack environment end-to-end at "
            "least one week before the actual hack. This will help ensure that "
            "all resources are properly configured and accessible. Report any "
            "issues to the provisioning team immediately."
        )

    # ──────────── Access Guide intro ────────────

    @staticmethod
    def _add_access_intro(doc: Document) -> None:
        doc.add_paragraph(
            "This document provides a comprehensive overview of the hackathon "
            "environment, including instructions on how to use the provisioned "
            "portals and services."
        )

    # ──────────── Dynamic Access Sections ────────────

    def _add_access_sections(self, doc: Document, state: Dict[str, Any]) -> None:
        users = state.get("users", [])
        all_licenses: List[str] = []
        for u in users:
            all_licenses.extend(u.get("licenses", []))
        all_licenses = list(set(all_licenses))

        if not all_licenses:
            doc.add_paragraph(
                "No specific licenses detected. Contact your admin for "
                "service access instructions."
            )
            return

        config = state.get("config", {})
        is_gcc = bool(config.get("isGcc", False))

        sections = get_sections_for_licenses(all_licenses, is_gcc=is_gcc)
        for _title, builder in sections:
            builder(doc, self._screenshots_dir)

    # ──────────── Support Contact ────────────

    @staticmethod
    def _add_support_contact(doc: Document, state: Dict[str, Any]) -> None:
        doc.add_heading("Support Contact", level=2)

        doc.add_paragraph(
            "A Teams channel will be set up for the purpose of issue reporting, "
            "escalation, and real-time support during the hackathon."
        )

        doc.add_paragraph("Members to include:")
        for member in ["Technical Team", "Support Team", "Local Microsoft team",
                        "Pre-event call attendees"]:
            doc.add_paragraph(member, style="List Bullet")

        # Response time table
        doc.add_paragraph("")
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Scenario"
        hdr[1].text = "Response Time"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

        for scenario, response in [
            ("Teams/email acknowledgment", "24 business hours"),
            ("Live ongoing event issue", "15 minutes"),
        ]:
            row = table.add_row()
            row.cells[0].text = scenario
            row.cells[1].text = response

        doc.add_paragraph("")
        doc.add_paragraph(
            'Urgent requests must use subject: "Immediate Support Required"'
        )

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Email Support: ")
        from .templates import _add_hyperlink
        _add_hyperlink(p, "mailto:cloudlabs-support@spektrasystems.com",
                       "cloudlabs-support@spektrasystems.com")

        p2 = doc.add_paragraph()
        p2.add_run("For Live Chat Support: ")
        _add_hyperlink(p2, "https://cloudlabs.ai/ms-support",
                       "https://cloudlabs.ai/ms-support")

        doc.add_paragraph("")
        p = doc.add_paragraph("This document was auto-generated by the Spektra hack setup system.")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = GRAY
